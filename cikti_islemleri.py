# cikti_islemleri.py - v4 (Dil ve Renklendirme Hataları Tamamen Giderildi - Kurşun Geçirmez Versiyon)

import pandas as pd
from openpyxl.styles import Font, Border, Side, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import RGBColor
import math

# ------------------ Yardımcı Fonksiyonlar ------------------

def format_tr_number(x: float, decimals: int = 2) -> str:
    if x is None or (isinstance(x, float) and (math.isnan(x) or math.isinf(x))):
        return ""
    try:
        return f"{x:,.{decimals}f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except (ValueError, TypeError):
        return ""

def is_percentage_key(k: str) -> bool:
    return any(kw in k for kw in ["Marjı", "Margin", "ROA", "ROE", "%"])

def apply_excel_header_style(ws, header_font, header_fill, border, is_pivot=False):
    # Sütun Başlıkları Stili
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
    # Pivot Tablo Stili (İlk Sütun)
    if is_pivot:
        for cell in ws['A']:
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border

def auto_fit_columns(ws):
    for column in ws.columns:
        max_length = 0
        column_letter = get_column_letter(column[0].column)
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = (max_length + 2) if max_length < 40 else 40
        ws.column_dimensions[column_letter].width = adjusted_width

# ------------------ Excel Çıktısı ------------------
def excel_cikti_al(finansal_veriler, oranlar, skorlar, detayli_analiz, dil, buffer) -> bool:
    try:
        writer = pd.ExcelWriter(buffer, engine="openpyxl")
        
        header_font = Font(bold=True, name='Calibri', size=11, color="FFFFFF")
        title_font = Font(bold=True, name='Calibri', size=14)
        border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
        header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")

        # --- Sayfa 1: Financial Data ---
        sheet_name_data = "Finansal Veriler" if dil == "TR" else "Financial Data"
        df_veriler = pd.DataFrame(finansal_veriler).set_index('yillar').T
        df_veriler.index = rows_tr if dil == "TR" else rows_en
        df_veriler.index.name = "Finansal Kalem" if dil == "TR" else "Financial Item"
        df_veriler.to_excel(writer, sheet_name=sheet_name_data)
        ws_data = writer.sheets[sheet_name_data]
        apply_excel_header_style(ws_data, header_font, header_fill, border, is_pivot=True)
        auto_fit_columns(ws_data)
        
        # --- Sayfa 2: Ratios ---
        sheet_name_ratios = "Oranlar" if dil == "TR" else "Ratios"
        oran_basliklari = oran_basliklari_tr if dil == "TR" else oran_basliklari_en
        df_oranlar = pd.DataFrame(oranlar, index=finansal_veriler['yillar']).T
        df_oranlar.index = [oran_basliklari.get(k, k) for k in df_oranlar.index]
        df_oranlar.index.name = "Oran" if dil == "TR" else "Ratio"
        df_oranlar.to_excel(writer, sheet_name=sheet_name_ratios)
        ws_ratios = writer.sheets[sheet_name_ratios]
        apply_excel_header_style(ws_ratios, header_font, header_fill, border, is_pivot=True)
        auto_fit_columns(ws_ratios)

        # --- Sayfa 3: Scores ---
        sheet_name_scores = "Skorlar" if dil == "TR" else "Scores"
        df_skorlar_raw = pd.DataFrame(skorlar).T.reset_index()
        if dil == "TR":
            df_skorlar_final = df_skorlar_raw[['index', 'toplam_skor', 'degerlendirme', 'olumlu_yonler', 'olumsuz_yonler']]
            df_skorlar_final.columns = ['Yıl', 'Toplam Skor', 'Değerlendirme', 'Olumlu Yönler', 'Olumsuz Yönler']
        else:
            df_skorlar_final = df_skorlar_raw[['index', 'toplam_skor', 'degerlendirme_en', 'olumlu_yonler_en', 'olumsuz_yonler_en']]
            df_skorlar_final.columns = ['Year', 'Total Score', 'Evaluation', 'Strengths', 'Weaknesses']
        df_skorlar_final.to_excel(writer, sheet_name=sheet_name_scores, index=False)
        ws_scores = writer.sheets[sheet_name_scores]
        apply_excel_header_style(ws_scores, header_font, header_fill, border)
        auto_fit_columns(ws_scores)
        
        # --- Sayfa 4: Detaylı Analiz ---
        sheet_name_analiz = "Detaylı Analiz" if dil == "TR" else "Detailed Analysis"
        ws_analiz = writer.book.create_sheet(title=sheet_name_analiz)
        start_row = 1
        for yil, veriler in detayli_analiz.items():
            if not veriler: continue
            
            title_text = f"{yil} " + ("Yılı Analizi" if dil == "TR" else "Year Analysis")
            ws_analiz.cell(row=start_row, column=1, value=title_text).font = title_font
            start_row += 2
            
            df_analiz = pd.DataFrame(veriler)
            
            for c_idx, col_name in enumerate(df_analiz.columns, 1):
                cell = ws_analiz.cell(row=start_row, column=c_idx, value=col_name)
                cell.font = header_font; cell.fill = header_fill; cell.border = border
            
            for r_idx, row_data in enumerate(df_analiz.itertuples(), start=start_row + 1):
                for c_idx, cell_value in enumerate(row_data[1:], 1):
                    cell = ws_analiz.cell(row=r_idx, column=c_idx, value=cell_value)
                    cell.border = border
                    # Düzeltilmiş Renklendirme Mantığı (Sütun sırasına göre)
                    if c_idx == 3: # 3. Sütun: Yıllık Değişim / YoY Change
                        if '▲' in str(cell_value): cell.font = Font(name='Calibri', size=11, color="008000")
                        elif '▼' in str(cell_value): cell.font = Font(name='Calibri', size=11, color="FF0000")
                    if c_idx == 5: # 5. Sütun: Durum / Status
                        if str(cell_value) in ["İyi", "Good"]: cell.font = Font(name='Calibri', size=11, color="008000")
                        elif str(cell_value) in ["Zayıf", "Weak"]: cell.font = Font(name='Calibri', size=11, color="FF0000")
                        elif str(cell_value) in ["Orta", "Average"]: cell.font = Font(name='Calibri', size=11, color="FFA500")

            start_row += len(df_analiz) + 3
        
        auto_fit_columns(ws_analiz)

        writer.close()
        return True
    except Exception as e:
        print(f"Excel çıktısı hatası: {e}")
        return False


# ------------------ Word Çıktısı ------------------
def word_cikti_al(finansal_veriler, oranlar, skorlar, detayli_analiz, dil, buffer) -> bool:
    try:
        doc = Document()
        oran_basliklari = oran_basliklari_tr if dil == "TR" else oran_basliklari_en
        yillar = finansal_veriler['yillar']

        doc.add_heading("Şirket Finansal Raporu" if dil == "TR" else "Company Financial Report", 0)

        # Bölüm 1: Finansal Veriler
        doc.add_heading("Finansal Veriler" if dil == "TR" else "Financial Data", level=1)
        # ... (Bu kısım doğru çalışıyordu, değişiklik yok)
        df_veriler = pd.DataFrame(finansal_veriler).set_index('yillar').T
        df_veriler.index = rows_tr if dil == "TR" else rows_en
        df_veriler.index.name = "Finansal Kalem" if dil == "TR" else "Financial Item"
        df_veriler = df_veriler.reset_index()
        table = doc.add_table(rows=1, cols=len(df_veriler.columns), style="Table Grid")
        for j, col_name in enumerate(df_veriler.columns): table.cell(0, j).text = str(col_name)
        for i, row in df_veriler.iterrows():
            row_cells = table.add_row().cells
            for j, value in enumerate(row):
                row_cells[j].text = format_tr_number(value, 2) if j > 0 else str(value)


        # Bölüm 2: Finansal Oranlar (Basit)
        doc.add_heading("Finansal Oranlar" if dil == "TR" else "Financial Ratios", level=1)
        # ... (Bu kısım doğru çalışıyordu, değişiklik yok)
        df_oranlar = pd.DataFrame(oranlar, index=yillar).T
        df_oranlar.index = [oran_basliklari.get(k, k) for k in df_oranlar.index]
        df_oranlar.index.name = "Oran" if dil == "TR" else "Ratio"
        df_oranlar = df_oranlar.reset_index()
        table = doc.add_table(rows=1, cols=len(df_oranlar.columns), style="Table Grid")
        for j, col_name in enumerate(df_oranlar.columns): table.cell(0, j).text = str(col_name)
        for i, row in df_oranlar.iterrows():
            row_cells = table.add_row().cells
            is_percent = is_percentage_key(row.iloc[0])
            for j, value in enumerate(row):
                if j > 0: row_cells[j].text = f"{value * 100:.2f}%" if is_percent and value is not None else f"{value:.2f}" if value is not None else ""
                else: row_cells[j].text = str(value)
        
        # Bölüm 3: Detaylı Oran Analizi
        doc.add_heading("Detaylı Oran Analizi" if dil == "TR" else "Detailed Ratio Analysis", level=1)
        for yil, veriler in detayli_analiz.items():
            if not veriler: continue
            doc.add_heading(f"{yil} " + ("Yılı Analizi" if dil == "TR" else "Year Analysis"), level=2)
            df = pd.DataFrame(veriler)
            table = doc.add_table(rows=1, cols=len(df.columns), style='Table Grid')
            for i, col_name in enumerate(df.columns): table.rows[0].cells[i].text = col_name
            for index, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, col_name in enumerate(df.columns):
                    cell_text = str(row[col_name])
                    run = row_cells[i].paragraphs[0].add_run(cell_text)
                    # Düzeltilmiş Renklendirme Mantığı (Sütun sırasına göre)
                    if i == 2: # 3. Sütun: Yıllık Değişim / YoY Change
                        if '▲' in cell_text: run.font.color.rgb = RGBColor(0, 128, 0)
                        elif '▼' in cell_text: run.font.color.rgb = RGBColor(255, 0, 0)
                    elif i == 4: # 5. Sütun: Durum / Status
                        if cell_text in ["İyi", "Good"]: run.font.color.rgb = RGBColor(0, 128, 0)
                        elif cell_text in ["Zayıf", "Weak"]: run.font.color.rgb = RGBColor(255, 0, 0)
                        elif cell_text in ["Orta", "Average"]: run.font.color.rgb = RGBColor(255, 165, 0)
            doc.add_paragraph()

        # Bölüm 4: Kredi Skoru Analizi
        doc.add_heading("Kredi Skoru Analizi" if dil == "TR" else "Credit Score Analysis", level=1)
        # ... (Bu kısım doğru çalışıyordu, değişiklik yok)
        for yil in yillar:
            skor_bilgisi = skorlar.get(yil)
            if not skor_bilgisi: continue
            doc.add_heading(f"{'Yıl' if dil == 'TR' else 'Year'}: {yil}", level=2)
            p = doc.add_paragraph(); p.add_run(f"{'Toplam Skor' if dil == 'TR' else 'Total Score'}: ").bold = True
            p.add_run(f"{skor_bilgisi.get('toplam_skor', 0):.0f}")
            p = doc.add_paragraph(); p.add_run(f"{'Değerlendirme' if dil == 'TR' else 'Evaluation'}: ").bold = True
            p.add_run(skor_bilgisi.get('degerlendirme' if dil == 'TR' else 'degerlendirme_en', ''))
            
            olumlu = skor_bilgisi.get('olumlu_yonler' if dil == "TR" else 'olumlu_yonler_en', [])
            olumsuz = skor_bilgisi.get('olumsuz_yonler' if dil == "TR" else 'olumsuz_yonler_en', [])
            doc.add_heading("Olumlu Yönler" if dil == "TR" else "Strengths", level=3)
            if olumlu:
                for item in olumlu: doc.add_paragraph(item, style='List Bullet')
            doc.add_heading("Olumsuz Yönler" if dil == "TR" else "Weaknesses", level=3)
            if olumsuz:
                for item in olumsuz: doc.add_paragraph(item, style='List Bullet')
            doc.add_paragraph()

        doc.save(buffer)
        return True
    except Exception as e:
        print(f"Word çıktısı hatası: {e}")
        return False

# Bu kısım programın çalışması için gerekli, dokunmuyoruz.
rows_tr = ["Dönen Varlıklar", "Duran Varlıklar", "Kısa Vadeli Yükümlülükler", "Uzun Vadeli Yükümlülükler", "Özkaynaklar", "Net Satışlar", "Satışların Maliyeti", "Faaliyet Giderleri", "Finansman Giderleri"]
rows_en = ["Current Assets", "Fixed Assets", "Short-Term Liabilities", "Long-Term Liabilities", "Equity", "Net Sales", "Cost of Goods Sold", "Operating Expenses", "Financial Expenses"]
oran_basliklari_tr = { "Cari Oran": "Cari Oran", "Likidite Oranı": "Likidite Oranı", "Finansal Kaldıraç": "Finansal Kaldıraç", "Borç / Özkaynak Oranı": "Borç / Özkaynak Oranı", "Net Kar Marjı": "Net Kar Marjı", "Brüt Kar Marjı": "Brüt Kar Marjı", "Faaliyet Kar Marjı": "Faaliyet Kar Marjı", "Aktif Karlılığı (ROA)": "Aktif Karlılığı (ROA)", "Özkaynak Karlılığı (ROE)": "Özkaynak Karlılığı (ROE)", "EBITDA": "EBITDA", "EBITDA Marjı": "EBITDA Marjı" }
oran_basliklari_en = { "Cari Oran": "Current Ratio", "Likidite Oranı": "Quick Ratio", "Finansal Kaldıraç": "Financial Leverage", "Borç / Özkaynak Oranı": "Debt to Equity Ratio", "Net Kar Marjı": "Net Profit Margin", "Brüt Kar Marjı": "Gross Profit Margin", "Faaliyet Kar Marjı": "Operating Profit Margin", "Aktif Karlılığı (ROA)": "Return on Assets (ROA)", "Özkaynak Karlılığı (ROE)": "Return on Equity (ROE)", "EBITDA": "EBITDA", "EBITDA Marjı": "EBITDA Margin" }